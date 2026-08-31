"""Build a .docx transcript for a video.

Options: which script(s) to include (roman / urdu / both) and whether to prefix
each line with its timestamp. Returns the .docx as bytes.

Roman lines that haven't been transliterated yet come through as "—"; export
Roman only after romanizing the video (the page offers that).
"""
import io


def _hhmmss(seconds):
    seconds = int(seconds or 0)
    h, rem = divmod(seconds, 3600)
    m, s = divmod(rem, 60)
    return f"{h}:{m:02d}:{s:02d}" if h else f"{m:02d}:{s:02d}"


def transcript_docx(video, script="both", timestamps=True):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    want_roman = script in ("roman", "both")
    want_urdu = script in ("urdu", "both")

    doc = Document()
    doc.add_heading(video["title"] or "Transcript", level=1)
    if video.get("source") == "user_upload":
        when = (video.get("uploaded_at") or "")[:10]
        doc.add_paragraph(f"User upload{' · uploaded ' + when if when else ''} · transcript in Urdu & Roman Urdu")
    elif video.get("youtube_url"):
        doc.add_paragraph(video["youtube_url"])

    for seg in video["segments"]:
        ts = _hhmmss(seg["start_time"])
        if want_roman:
            roman = seg.get("roman_text") or "—"
            p = doc.add_paragraph()
            if timestamps:
                run = p.add_run(f"[{ts}] ")
                run.bold = True
            p.add_run(roman)
        if want_urdu:
            urdu = seg.get("urdu_text") or ""
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Urdu reads right-to-left
            if timestamps and not want_roman:
                run = p.add_run(f"[{ts}] ")
                run.bold = True
            p.add_run(urdu)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def transcript_csv(video):
    """The transcript as CSV bytes: timestamp, urdu, roman — one row per segment.
    UTF-8 with BOM so Excel renders the Urdu script correctly."""
    import csv

    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["timestamp", "urdu", "roman"])
    for seg in video["segments"]:
        w.writerow([_hhmmss(seg["start_time"]), seg.get("urdu_text") or "", seg.get("roman_text") or ""])
    return buf.getvalue().encode("utf-8-sig")
